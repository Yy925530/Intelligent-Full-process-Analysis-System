"""
销售数据全流程智能分析系统
================================

设计思路（依据《设计思路简述》）：
1. 基础层：数据清洗与质量评估
   - 统一处理4份数据的缺失值、异常值、格式问题、逻辑矛盾、重复值
   - 字段自适应适配（月饼/巧克力/家具服装/含用户数据）

2. 路径1：描述性分析（趋势/分布/对比）
   - 趋势分析：按销售日期做月度/季度销售额、销量趋势，识别淡旺季
   - 分布分析：价格、销量、销售额的分布情况
   - 对比分析：地区维度、产品维度、用户维度、渠道维度、销售人员维度

3. 路径2：相关性与影响因素分析
   - 变量相关性分析：计算价格、销量、折扣、年龄等变量与销售额的相关性
   - 单变量影响分析：价格对销量的影响、折扣对销售额的影响、年龄/性别对消费的影响

4. 路径3：多模型对比分析
   - 多元线性回归、决策树回归、支持向量回归、随机森林回归
   - 对比R²、MAE、RMSE，自动选择最优模型

5. 路径4：变量交互分析
   - 地区×产品类别、年龄×产品类别、渠道×支付方式、价格×折扣

6. 报告导出：自动生成企业级经营分析报告（WORD）
"""
import streamlit as st
import pandas as pd
import numpy as np
from scipy import stats
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
import matplotlib.pyplot as plt
import seaborn as sns
import warnings
import io
from datetime import datetime
warnings.filterwarnings('ignore')
# 页面全局配置
st.set_page_config(page_title="销售数据全流程智能分析系统", layout="wide")
st.title("销售数据全流程智能分析系统")
st.markdown("**数据清洗 → 异常检测 → 可视化分析 → 多模型对比 → 报告生成**")
# matplotlib全局固定字体配置，所有绘图统一生效
import matplotlib
# Agg必须放在import plt最前面，解决云端空白
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import platform
# Windows：黑体；Linux云端：Noto Sans CJK SC（系统预装中文，不用DejaVu）
if platform.system() == "Linux":
    plt.rcParams["font.sans-serif"] = ["Noto Sans CJK SC"]
else:
    plt.rcParams["font.sans-serif"] = ["SimHei","Microsoft YaHei"]
plt.rcParams["axes.unicode_minus"] = False
# 刷新mat字体缓存，改字体立刻生效
from matplotlib.font_manager import _rebuild
_rebuild()

# ==================== 辅助函数定义 ====================

def clean_price_format(price_value):
    """清洗价格格式：去除单位、处理科学计数法"""
    if pd.isna(price_value):
        return np.nan
    price_str = str(price_value).strip()
    price_str = price_str.replace("元", "").replace("￥", "").replace("$", "")
    if 'e' in price_str.lower():
        try:
            return float(price_str)
        except:
            pass
    try:
        return float(price_str)
    except:
        return np.nan


def clean_region_format(region_value):
    """清洗地区格式：统一大小写、展开简称"""
    if pd.isna(region_value):
        return np.nan
    region_str = str(region_value).strip()
    if region_str.lower() in ['north', 'south', 'east', 'west', 'uk', 'india', 'australia']:
        region_str = region_str.title()
    region_short_map = {'渝': '重庆', '京': '北京', '沪': '上海', '穗': '广州', '深': '深圳'}
    for short, full in region_short_map.items():
        if short in region_str:
            region_str = region_str.replace(short, full)
    return region_str


def clean_product_format(product_value):
    """清洗产品名称格式：去除首尾空格、特殊符号"""
    if pd.isna(product_value):
        return np.nan
    product_str = str(product_value).strip()
    product_str = product_str.replace("###", "").strip()
    return product_str if product_str else np.nan


def extract_moon_flavor(product_name):
    """从月饼产品名称中提取口味"""
    if pd.isna(product_name):
        return "未知"
    p = str(product_name).lower()
    if '流心奶黄' in p or ('流心' in p and '奶黄' in p):
        return '流心奶黄'
    if '五仁' in p or '伍仁' in p:
        return '五仁'
    if '莲蓉' in p:
        return '莲蓉'
    if '蛋黄' in p:
        return '蛋黄'
    if '豆沙' in p:
        return '豆沙'
    if '冰皮' in p:
        return '冰皮'
    if '云腿' in p or '火腿' in p:
        return '云腿'
    if '巧克力' in p:
        return '巧克力'
    if '水果' in p:
        return '水果'
    if '椰蓉' in p:
        return '椰蓉'
    if '枣泥' in p:
        return '枣泥'
    return '其他'


def extract_choco_flavor(product_name):
    """从巧克力产品名称中提取口味"""
    if pd.isna(product_name):
        return "未知"
    p = str(product_name).lower()
    if 'dark' in p:
        return '黑巧克力(Dark)'
    if 'milk' in p:
        return '牛奶巧克力(Milk)'
    if 'white' in p:
        return '白巧克力(White)'
    if 'peanut' in p:
        return '花生酱'
    if 'almond' in p:
        return '杏仁'
    if 'raspberry' in p:
        return '覆盆子'
    if 'orange' in p:
        return '橙子'
    if 'mint' in p:
        return '薄荷'
    if 'caramel' in p:
        return '焦糖'
    if 'eclairs' in p:
        return '闪电泡芙'
    if 'manuka' in p:
        return '麦卢卡蜂蜜'
    return '其他'


def find_column(df, keywords, return_first=True):
    """模糊查找包含指定关键词的列"""
    matched = []
    for col in df.columns:
        for kw in keywords:
            if kw in col:
                matched.append(col)
                if return_first:
                    return col
    return matched if matched else None


def clean_group_data(series):
    """清洗分组用的列"""
    return series.astype(str).str.strip().str.lower().replace("", np.nan).dropna()


def standardize_gender(gender):
    """标准化性别：统一为 Male / Female 格式"""
    if pd.isna(gender):
        return np.nan
    g = str(gender).strip()
    if g.lower() in ['male', 'm']:
        return 'Male'
    elif g.lower() in ['female', 'f']:
        return 'Female'
    else:
        return g


def encode_city_level(region_name):
    """地区有序编码：一线=5, 新一线=4, 二线=3, 三线=2, 其他=1"""
    if pd.isna(region_name):
        return 1
    r = str(region_name)
    if any(c in r for c in ['北京', '上海', '广州', '深圳']):
        return 5
    if any(c in r for c in ['成都', '杭州', '重庆', '武汉', '西安', '天津', '南京', '郑州', '长沙', '沈阳', '青岛', '苏州', '东莞', '佛山']):
        return 4
    if any(c in r for c in ['昆明', '哈尔滨', '长春', '南昌', '南宁', '合肥', '太原', '石家庄', '乌鲁木齐', '贵阳', '厦门', '福州']):
        return 3
    if any(c in r for c in ['嘉兴', '无锡', '金华', '湖州', '绍兴', '温州', '惠州', '中山', '珠海', '镇江', '扬州', '南通']):
        return 2
    return 1


def encode_product_tier(price_series, product_series):
    """产品类别有序编码：按平均价格从低到高赋1-5"""
    if product_series is None or price_series is None:
        return {}
    df_temp = pd.DataFrame({'产品类别': product_series, '价格': price_series}).dropna()
    if df_temp.empty:
        return {}
    avg_price = df_temp.groupby('产品类别')['价格'].mean().sort_values()
    n = len(avg_price)
    if n <= 1:
        return {cat: 3 for cat in avg_price.index}
    breaks = np.percentile(range(n), [20, 40, 60, 80])
    tier_map = {}
    for i, cat in enumerate(avg_price.index):
        if i < breaks[0]:
            tier_map[cat] = 1
        elif i < breaks[1]:
            tier_map[cat] = 2
        elif i < breaks[2]:
            tier_map[cat] = 3
        elif i < breaks[3]:
            tier_map[cat] = 4
        else:
            tier_map[cat] = 5
    return tier_map


# ==================== 数据上传与读取 ====================
st.divider()
st.subheader("数据上传")

upload_file = st.file_uploader("请拖入或上传数据文件（支持Excel/CSV）", type=["xlsx", "csv"])

df = None
raw_df = None
original_rows = 0

if upload_file:
    if upload_file.name.endswith(".csv"):
        df = pd.read_csv(upload_file)
    else:
        df = pd.read_excel(upload_file)
    raw_df = df.copy()
    original_rows = len(df)
    st.success(f"✅ 数据读取成功！原始数据：共 {original_rows} 行，{len(df.columns)} 列")
    st.write("**数据预览（前5行）**：")
    st.dataframe(df.head())
else:
    st.info("请上传数据文件以开始分析")
    st.stop()

# ==================== 全类型异常值检测 ====================
st.divider()
st.subheader("全类型异常值检测")

# 初始化异常计数器
dup_num = df.duplicated().sum()
miss_num = df.isnull().sum().sum()
text_in_num = 0
num_in_text = 0
special_symbol = 0
negative_err = 0
extreme_outlier = 0
gender_err = 0
date_err = 0
zero_price_non_sales = 0
zero_sales_non_amount = 0
amount_deviation_err = 0
age_outlier_err = 0
part_dup_num = 0

# 逐列扫描基础异常
for col in df.columns:
    col_data = df[col].astype(str).str.strip()
    special_symbol += col_data.str.contains(r'[,\，元￥$岁#@!]').sum()
    special_symbol += col_data.str.contains(r'\s{2,}').sum()
    
    if any(key in col for key in ["销售额", "价格", "销量", "单位成本", "折扣", "单价", "金额", "年龄"]):
        temp_num = pd.to_numeric(df[col], errors="coerce")
        text_in_num += temp_num.isna().sum()
        negative_err += (temp_num.dropna() < 0).sum()
        
        if "年龄" in col:
            age_outlier_err += ((temp_num < 0) | (temp_num > 120)).sum()
        
        if not temp_num.dropna().empty:
            q1 = temp_num.quantile(0.25)
            q3 = temp_num.quantile(0.75)
            iqr = q3 - q1
            extreme_outlier += ((temp_num < q1 - 1.5*iqr) | (temp_num > q3 + 1.5*iqr)).sum()
    
    if any(key in col for key in ["产品", "地区", "销售人员", "产品类别", "支付方式", "顾客类型", "销售渠道"]):
        num_in_text += col_data.str.contains(r"[0-9]").sum()

# 逻辑异常检测
if "价格" in df.columns and "销量" in df.columns:
    price_num = pd.to_numeric(df["价格"], errors="coerce")
    sales_num = pd.to_numeric(df["销量"], errors="coerce")
    zero_price_non_sales = df[(price_num == 0) & (sales_num > 0)].shape[0]

if "销量" in df.columns and "销售额" in df.columns:
    sales_num = pd.to_numeric(df["销量"], errors="coerce")
    amount_num = pd.to_numeric(df["销售额"], errors="coerce")
    zero_sales_non_amount = df[(sales_num == 0) & (amount_num > 0)].shape[0]

if "价格" in df.columns and "销量" in df.columns and "销售额" in df.columns:
    price_num = pd.to_numeric(df["价格"], errors="coerce")
    sales_num = pd.to_numeric(df["销量"], errors="coerce")
    amount_num = pd.to_numeric(df["销售额"], errors="coerce")
    valid_mask = price_num.notna() & sales_num.notna() & amount_num.notna()
    if valid_mask.sum() > 0:
        df_valid = df[valid_mask].copy()
        df_valid["理论销售额"] = price_num[valid_mask] * sales_num[valid_mask]
        deviation_mask = (abs(amount_num[valid_mask] - df_valid["理论销售额"]) / 
                         df_valid["理论销售额"].replace(0, np.nan) > 0.2)
        amount_deviation_err = deviation_mask.sum()

if "性别" in df.columns:
    valid_gender_base = ["male", "female", "m", "f", "男", "女", "男生", "女生", "男性", "女性"]
    gender_err = df[~df["性别"].astype(str).str.lower().isin(valid_gender_base)].shape[0]

date_col = find_column(df, ["销售日期", "日期", "date"])
if date_col:
    temp_date = pd.to_datetime(df[date_col], errors="coerce")
    date_err = temp_date.isna().sum()

key_cols = [c for c in ["产品", "价格", "年龄", "性别"] if c in df.columns]
if key_cols:
    part_dup_num = df.duplicated(subset=key_cols).sum()

# 展示检测结果
st.markdown("### 异常检测统计")
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("完全重复数据", dup_num)
    st.metric("核心字段部分重复", part_dup_num)
    st.metric("全局缺失值总数", miss_num)
with col2:
    st.metric("数值列混入文本", text_in_num)
    st.metric("文本列混入数字", num_in_text)
    st.metric("特殊符号/格式脏数据", special_symbol)
with col3:
    st.metric("不合理负值", negative_err)
    st.metric("年龄超范围", age_outlier_err)
    st.metric("极端离群值", extreme_outlier)

st.markdown("---")
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("非法性别值", gender_err)
    st.metric("日期格式错误", date_err)
with col2:
    st.metric("价格为0但销量>0", zero_price_non_sales)
    st.metric("销量为0但销售额>0", zero_sales_non_amount)
with col3:
    st.metric("销售额计算偏差", amount_deviation_err)

# ==================== 数据清洗 ====================
st.divider()
st.subheader("数据清洗")
df_clean = df.copy()

# 重复值清洗
before_dup = len(df_clean)
df_clean = df_clean.drop_duplicates()
st.write(f"✅ 已删除完全重复行：{before_dup - len(df_clean)} 条")

if key_cols:
    before_part = len(df_clean)
    df_clean = df_clean.drop_duplicates(subset=key_cols, keep="first")
    st.write(f"✅ 已删除部分重复行：{before_part - len(df_clean)} 条")

# 格式清洗
for col in df_clean.columns:
    if any(key in col for key in ["销售额", "价格", "销量", "单位成本", "折扣", "年龄", "单价", "金额"]):
        df_clean[col] = df_clean[col].astype(str).str.replace(r"[^0-9.\-]", "", regex=True)
        df_clean[col] = pd.to_numeric(df_clean[col], errors="coerce")

if "产品" in df_clean.columns:
    df_clean["产品"] = df_clean["产品"].apply(clean_product_format)

if "地区" in df_clean.columns:
    df_clean["地区"] = df_clean["地区"].apply(clean_region_format)

if date_col and date_col in df_clean.columns:
    df_clean[date_col] = pd.to_datetime(df_clean[date_col], errors="coerce")

if "性别" in df_clean.columns:
    df_clean["性别"] = df_clean["性别"].apply(standardize_gender)

# 缺失值处理
key_drop_cols = [c for c in ["产品", "价格", "地区"] if c in df_clean.columns]
if key_drop_cols:
    before_drop = len(df_clean)
    df_clean = df_clean.dropna(subset=key_drop_cols)
    st.write(f"✅ 已删除关键字段缺失行：{before_drop - len(df_clean)} 条")

numeric_cols = df_clean.select_dtypes(include=[np.number]).columns.tolist()
for col in numeric_cols:
    if not df_clean[col].dropna().empty:
        df_clean[col] = df_clean[col].fillna(df_clean[col].median())

for col in ["地区", "性别"]:
    if col in df_clean.columns:
        df_clean[col] = df_clean[col].fillna("未知")

# 逻辑错误修复
if "销量" in df_clean.columns and "销售额" in df_clean.columns:
    df_clean["销量"] = pd.to_numeric(df_clean["销量"], errors="coerce")
    df_clean["销售额"] = pd.to_numeric(df_clean["销售额"], errors="coerce")
    df_clean.loc[(df_clean["销量"] == 0) & (df_clean["销售额"] > 0), "销售额"] = 0

if "价格" in df_clean.columns and "销量" in df_clean.columns and "产品" in df_clean.columns:
    df_clean["价格"] = pd.to_numeric(df_clean["价格"], errors="coerce")
    df_clean["销量"] = pd.to_numeric(df_clean["销量"], errors="coerce")
    zero_price_mask = (df_clean["价格"] == 0) & (df_clean["销量"] > 0)
    if zero_price_mask.any():
        avg_price = df_clean.groupby("产品")["价格"].transform(lambda x: x[x>0].mean() if x[x>0].any() else 0)
        df_clean.loc[zero_price_mask, "价格"] = avg_price

if "价格" in df_clean.columns and "销量" in df_clean.columns and "销售额" in df_clean.columns:
    df_clean["理论销售额"] = df_clean["价格"] * df_clean["销量"]
    deviation_mask = (abs(df_clean["销售额"] - df_clean["理论销售额"]) / 
                     df_clean["理论销售额"].replace(0, np.nan) > 0.2)
    df_clean.loc[deviation_mask, "销售额"] = df_clean.loc[deviation_mask, "理论销售额"]
    df_clean = df_clean.drop(columns=["理论销售额"])

# 异常值剔除
for col in ["价格", "销量"]:
    if col in df_clean.columns:
        df_clean[col] = pd.to_numeric(df_clean[col], errors="coerce")
        df_clean = df_clean[df_clean[col] >= 0]

if "年龄" in df_clean.columns:
    df_clean["年龄"] = pd.to_numeric(df_clean["年龄"], errors="coerce")
    df_clean = df_clean[(df_clean["年龄"] >= 0) & (df_clean["年龄"] <= 120)]

if len(df_clean) > 100 and numeric_cols:
    valid_numeric = df_clean[numeric_cols].dropna()
    if not valid_numeric.empty:
        z_scores = np.abs(stats.zscore(valid_numeric))
        outlier_mask = pd.Series((z_scores < 3).all(axis=1), index=valid_numeric.index)
        outlier_mask = outlier_mask.reindex(df_clean.index, fill_value=True)
        df_clean = df_clean[outlier_mask]
        st.write("✅ 极端偏离的异常值已剔除")

final_rows = len(df_clean)
st.success(f"✅ 全部清洗完成，最终有效数据：{final_rows} 行")

# ==================== 有序编码说明（文献支持） ====================
st.divider()
st.subheader("📘 有序编码说明（文献支持）")
with st.expander("点击查看编码规则与文献依据", expanded=True):
    st.markdown("""
    **一、为什么采用有序编码？**  
    根据张文彤《SPSS统计分析高级教程》、吴喜之《统计学：从数据到结论》以及于斌斌(2022)《地理科学》等权威文献，对于高基数无序分类变量（如地区、产品类别），直接使用哑变量会导致维度爆炸、模型解释困难。  
    可行做法：依据**现实经济意义**将无序分类重构为**有序分类**，然后赋值为1/2/3/4/5等整数，作为连续变量纳入线性回归。该法统计合法，且被大量实证论文采用。

    **二、本系统编码规则**  
    **1. 地区编码（基于城市等级）**  
    - 一线城市 → 5  
    - 新一线城市 → 4  
    - 二线城市 → 3  
    - 三线城市 → 2  
    - 其他/境外 → 1  

    **2. 产品类别编码（基于平均价格档次）**  
    自动计算每个类别的平均价格，按价格升序等频划分为5档：  
    - 最低20% → 1（低端）  
    - 20%-40% → 2（中低）  
    - 40%-60% → 3（中端）  
    - 60%-80% → 4（中高）  
    - 80%-100% → 5（高端）  
    """)

# ==================== 先统一查找所有可能的列 ====================
region_col = find_column(df_clean, ["地区", "省份", "城市"])
product_col = find_column(df_clean, ["产品", "商品", "品名"])
category_col = find_column(df_clean, ["产品类别", "商品类别", "品类"])
gender_col = find_column(df_clean, ["性别"])
age_col = find_column(df_clean, ["年龄", "岁数"])
salesperson_col = find_column(df_clean, ["销售人员", "销售代表", "业务员"])
channel_col = find_column(df_clean, ["销售渠道", "渠道"])
pay_col = find_column(df_clean, ["支付方式", "付款方式"])
customer_type_col = find_column(df_clean, ["顾客类型", "客户类型"])

# 执行编码（如果存在相应列）
if region_col:
    df_clean["地区_编码"] = df_clean[region_col].apply(encode_city_level)

if category_col and "价格" in df_clean.columns:
    tier_map = encode_product_tier(df_clean["价格"], df_clean[category_col])
    if tier_map:
        df_clean["产品类别_编码"] = df_clean[category_col].map(tier_map)
# ==================== 路径1：描述性与趋势分析 ====================
st.divider()
st.subheader("路径1：描述性与趋势分析")

# 核心指标
total_sales = df_clean["销售额"].sum() if "销售额" in df_clean.columns else 0
avg_sales = df_clean["销售额"].mean() if "销售额" in df_clean.columns else 0
total_orders = len(df_clean)

st.markdown("### 核心销售指标概览")
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("总销售额", f"{total_sales:,.2f} 元")
with col2:
    st.metric("平均单笔销售额", f"{avg_sales:,.2f} 元")
with col3:
    st.metric("总订单数", f"{total_orders} 笔")
# 先看图数据，再结论决策
st.write("**数据解读：** 整体订单总量、总营收与单客均价构成整体销售基准，可衡量整体消费水平。")
st.info("**经营决策：** 以单笔均价作为产品定价基准，高于均价定位中高端产品，低于均价做引流特价品。")

# 趋势分析（时间序列双图）
if date_col and date_col in df_clean.columns:
    st.markdown("### 销售趋势分析（时间序列）")
    df_date_valid = df_clean[df_clean[date_col].notna()].copy()
    if not df_date_valid.empty:
        df_date_valid[date_col] = pd.to_datetime(df_date_valid[date_col], errors="coerce")
        df_date_valid = df_date_valid[df_date_valid[date_col].notna()]
        
        if not df_date_valid.empty:
            df_date_valid['销售月份'] = df_date_valid[date_col].dt.to_period('M')
            monthly_sales = df_date_valid.groupby('销售月份')["销售额"].sum()
            monthly_quantity = df_date_valid.groupby('销售月份')["销量"].sum() if "销量" in df_date_valid.columns else None
            
            fig, axes = plt.subplots(1, 2, figsize=(10, 3), dpi=120)
            monthly_sales.plot(ax=axes[0], marker='o', color="#2076b4", linewidth=2)
            axes[0].set_title("月度销售额趋势", fontsize=10)
            axes[0].set_xlabel("月份", fontsize=8)
            axes[0].set_ylabel("销售额（元）", fontsize=8)
            axes[0].grid(True, alpha=0.3)
            
            if monthly_quantity is not None:
                monthly_quantity.plot(ax=axes[1], marker='s', color='#ff7f0e', linewidth=2)
                axes[1].set_title("月度销量趋势", fontsize=10)
                axes[1].set_xlabel("月份", fontsize=8)
                axes[1].set_ylabel("销量", fontsize=8)
                axes[1].grid(True, alpha=0.3)
            
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            
            if len(monthly_sales) >= 3:
                peak_month = monthly_sales.idxmax()
                trough_month = monthly_sales.idxmin()
                st.write(f"📈 销售旺季：{peak_month}，销售淡季：{trough_month}")
                st.write("**数据解读：** 月度销售额与销量走势基本同步，存在明显淡旺季，峰值月份营收远高于低谷月份。")
                st.info("**经营决策：** 旺季{peak_month}提前备货、加大广告投放；淡季{trough_month}开展满减、捆绑促销提升销量。")

# 价格&销量分布直方图
st.markdown("### 价格与销量分布分析")
col1, col2 = st.columns(2)
with col1:
    if "价格" in df_clean.columns and not df_clean["价格"].dropna().empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        df_clean["价格"].hist(bins=20, color="#2ca02c", edgecolor='black', alpha=0.7)
        ax.set_title("价格分布直方图", fontsize=10)
        ax.set_xlabel("价格（元）", fontsize=8)
        ax.set_ylabel("频数", fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        price_mean = df_clean["价格"].mean()
        st.write(f"**数据解读：** 商品均价{price_mean:.2f}元，直方图柱子集中区间为市场接受主力价位，两端为低价/高价小众产品。")
        st.info("**经营决策：** 主力价位产品作为店铺主推，少量高低价产品完善产品结构。")
with col2:
    if "销量" in df_clean.columns and not df_clean["销量"].dropna().empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        df_clean["销量"].hist(bins=20, color="#ff7f0e", edgecolor='black', alpha=0.7)
        ax.set_title("销量分布直方图", fontsize=10)
        ax.set_xlabel("销量", fontsize=8)
        ax.set_ylabel("频数", fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        sale_mean = df_clean["销量"].mean()
        st.write(f"**数据解读：** 单订单平均采购量{sale_mean:.2f}件，多数订单小批量采购，大额批量订单占比偏低。")
        st.info("**经营决策：** 设计多件组合优惠，引导用户提高单次采购数量。")

# ==================== 多维度对比分析 ====================
st.markdown("### 多维度对比分析")
# 地区销售额柱状
if region_col:
    df_clean["地区_清洗"] = clean_group_data(df_clean[region_col])
    region_sales = df_clean.groupby("地区_清洗")["销售额"].sum().sort_values(ascending=False)
    region_sales = region_sales[region_sales > 0]
    
    col1, col2 = st.columns(2)
    with col1:
        if not region_sales.empty:
            fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
            region_sales.head(10).plot(kind="bar", color="#9467bd", ax=ax)
            ax.set_title("各地区销售额对比（前10）", fontsize=10)
            ax.set_xlabel("地区", fontsize=8)
            ax.set_ylabel("销售额（元）", fontsize=8)
            plt.xticks(rotation=45, ha="right", fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            st.write("**数据解读：** 区域营收差距悬殊，头部地区消费体量遥遥领先，尾部地区市场潜力不足。")
            st.info("**经营决策：** 深耕头部热门区域，增设线下/推广资源；低销地区调研消费习惯，针对性选品拓客。")
    with col2:
        if len(region_sales) >= 3:
            top3 = region_sales.head(3)
            bottom3 = region_sales.tail(3)
            st.write("**销售额TOP3地区**")
            for i, (region, sales) in enumerate(top3.items(), 1):
                st.write(f"{i}. {region}: {sales:,.0f}元")
            st.write("**销售额末位3地区**")
            for i, (region, sales) in enumerate(bottom3.items(), 1):
                st.write(f"{i}. {region}: {sales:,.0f}元")
            top_region = top3.index[0]
            st.success(f"📌 **购买力最强的地区：{top_region}**，贡献了 {top3.iloc[0]/region_sales.sum()*100:.1f}% 的销售额。")

# 单品TOP10柱状
if product_col:
    df_clean["产品_清洗"] = clean_group_data(df_clean[product_col])
    product_sales = df_clean.groupby("产品_清洗")["销售额"].sum().sort_values(ascending=False).head(10)
    if not product_sales.empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        product_sales.plot(kind="bar", color="#8c564b", ax=ax)
        ax.set_title("产品销售额Top10",fontsize=10)
        ax.set_xlabel("产品",fontsize=8)
        ax.set_ylabel("销售额（元）",fontsize=8)
        plt.xticks(rotation=45, ha="right",fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        top_product = product_sales.index[0]
        st.success(f"📌 **最畅销产品：{top_product}**，销售额 {product_sales.iloc[0]:,.0f}元。")
        st.write("**数据解读：** 头部爆款单品贡献大量营收，后面单品销量快速下滑，单品销售集中度高。")
        st.info("**经营决策：** 爆款作为引流主力，加大库存备货；滞销单品优化定价或淘汰下架。")

# 品类&性别双饼图
if category_col or gender_col:
    col1, col2 = st.columns(2)
    with col1:
        if category_col:
            df_clean["类别_清洗"] = clean_group_data(df_clean[category_col])
            category_sales = df_clean.groupby("类别_清洗")["销售额"].sum()
            category_sales = category_sales[category_sales > 0]
            if not category_sales.empty:
                fig, ax = plt.subplots(figsize=(3.2, 3), dpi=120)
                category_sales.plot(kind="pie", autopct="%1.1f%%", ax=ax, wedgeprops=dict(width=0.7))
                ax.set_title("各产品类别销售额占比",fontsize=10)
                ax.set_ylabel("")
                plt.tight_layout(pad=0.3)
                st.pyplot(fig, use_container_width=False)
                top_cat = category_sales.idxmax()
                st.success(f"📌 **最畅销产品类别：{top_cat}**，占比 {category_sales.max()/category_sales.sum()*100:.1f}%")
                st.write("**数据解读：** 单一品类占比过高，品类结构不均衡，弱势品类营收体量偏小。")
                st.info("**经营决策：** 巩固王牌品类优势，扩充弱势品类SKU，平衡营收结构。")
                if 'Clothing' in category_sales.index and 'Books' in category_sales.index:
                    clothing_sales = category_sales['Clothing']
                    books_sales = category_sales['Books']
                    if clothing_sales > books_sales:
                        st.write(f"   - 服装(Clothing)销售额 ({clothing_sales:,.0f}元) 高于图书(Books) ({books_sales:,.0f}元)。")
                    else:
                        st.write(f"   - 图书(Books)销售额 ({books_sales:,.0f}元) 高于服装(Clothing) ({clothing_sales:,.0f}元)。")
    with col2:
        if gender_col:
            gender_sales = df_clean[df_clean["性别"].isin(["Male", "Female"])].groupby("性别")["销售额"].sum()
            if not gender_sales.empty:
                fig, ax = plt.subplots(figsize=(3.2, 3), dpi=120)
                gender_sales.plot(kind="pie", autopct="%1.1f%%", ax=ax, wedgeprops=dict(width=0.7))
                ax.set_title("不同性别消费占比",fontsize=10)
                ax.set_ylabel("")
                plt.tight_layout(pad=0.3)
                st.pyplot(fig, use_container_width=False)
                if gender_sales['Female'] > gender_sales['Male']:
                    st.success(f"📌 **女性消费者贡献更高**，占比 {gender_sales['Female']/gender_sales.sum()*100:.1f}%。")
                    st.write("**数据解读：** 女性是消费主力人群，消费贡献远超男性客群。")
                    st.info("**经营决策：** 产品选型、活动营销围绕女性消费偏好设计。")
                else:
                    st.success(f"📌 **男性消费者贡献更高**，占比 {gender_sales['Male']/gender_sales.sum()*100:.1f}%。")
                    st.write("**数据解读：** 男性为核心消费群体。")
                    st.info("**经营决策：** 产品开发与营销侧重男性需求。")

# 口味柱状
if product_col:
    sample_prod = str(df_clean[product_col].iloc[0]) if len(df_clean) > 0 else ""
    is_moon = "月饼" in sample_prod
    is_choco = "巧克力" in sample_prod or "Choco" in sample_prod or "Dark" in sample_prod
    if is_moon or is_choco:
        st.markdown("### 产品口味/类型分析")
        if is_moon:
            df_clean["口味"] = df_clean[product_col].apply(extract_moon_flavor)
        else:
            df_clean["口味"] = df_clean[product_col].apply(extract_choco_flavor)
        flavor_sales = df_clean.groupby("口味")["销售额"].sum().sort_values(ascending=False)
        if not flavor_sales.empty:
            fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
            flavor_sales.head(10).plot(kind="bar", color="coral", ax=ax)
            ax.set_title("各口味/类型销售额排名",fontsize=10)
            ax.set_xlabel("口味/类型",fontsize=8)
            ax.set_ylabel("销售额（元）",fontsize=8)
            plt.xticks(rotation=75, ha="right", fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            top_flavor = flavor_sales.index[0]
            st.success(f"📌 **最受欢迎的口味：{top_flavor}**，销售额占比 {flavor_sales.iloc[0]/flavor_sales.sum()*100:.1f}%")
            st.write("**数据解读：** 头部口味市场认可度高，小众口味市场需求有限。")
            st.info("**经营决策：** 爆款口味加大产能，小众口味小批量备货减少库存积压。")

# 年龄分层柱状
if age_col:
    st.markdown("### 不同年龄层消费分布")
    df_clean["年龄_数值"] = pd.to_numeric(df_clean[age_col], errors="coerce")
    df_age_valid = df_clean[(df_clean["年龄_数值"] >= 0) & (df_clean["年龄_数值"] <= 120)]
    if not df_age_valid.empty:
        df_age_valid["年龄分组"] = pd.cut(df_age_valid["年龄_数值"], bins=[0,25,35,45,60,120],
                                          labels=["18-25岁", "26-35岁", "36-45岁", "46-60岁", "60+岁"])
        age_sales = df_age_valid.groupby("年龄分组")["销售额"].sum()
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        age_sales.plot(kind="bar", color="#e377c2", ax=ax)
        ax.set_title("各年龄层销售额对比",fontsize=10)
        ax.set_xlabel("年龄分组",fontsize=8)
        ax.set_ylabel("销售额（元）",fontsize=8)
        plt.xticks(fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        top_age = age_sales.idxmax()
        st.success(f"📌 **消费主力年龄段：{top_age}**，贡献销售额 {age_sales.max():,.0f}元。")
        st.write("**数据解读：** {top_age}年龄段消费能力最强，其余年龄段消费体量依次递减。")
        st.info("**经营决策：** 资源倾斜主力年龄段产品，针对其他年龄段开发适配刚需产品。")

# 销售人员业绩
if salesperson_col:
    st.markdown("### 销售人员业绩Top10")
    df_clean["销售_清洗"] = clean_group_data(df_clean[salesperson_col])
    salesperson_sales = df_clean.groupby("销售_清洗")["销售额"].sum().sort_values(ascending=False).head(10)
    if not salesperson_sales.empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        salesperson_sales.plot(kind="bar", color="#d62728", ax=ax)
        ax.set_title("销售人员业绩Top10",fontsize=10)
        ax.set_xlabel("销售人员",fontsize=8)
        ax.set_ylabel("销售额（元）",fontsize=8)
        plt.xticks(rotation=45, ha="right", fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        top_sp = salesperson_sales.index[0]
        st.success(f"📌 **最佳销售人员：{top_sp}**，业绩 {salesperson_sales.iloc[0]:,.0f}元，占比 {salesperson_sales.iloc[0]/salesperson_sales.sum()*100:.1f}%。")
        st.write("**数据解读：** 销冠业绩远超团队平均水平，内部人员业绩分化明显。")
        st.info("**经营决策：** 销冠分享销售话术与拓客经验，培训落后员工，设置阶梯提成激励全员。")

# 渠道营收
if channel_col:
    st.markdown("### 销售渠道销售额对比")
    df_clean["渠道_清洗"] = clean_group_data(df_clean[channel_col])
    channel_sales = df_clean.groupby("渠道_清洗")["销售额"].sum()
    channel_sales = channel_sales[channel_sales > 0]
    if not channel_sales.empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        channel_sales.plot(kind="bar", color="#4CAF50", ax=ax)
        ax.set_title("各销售渠道销售额对比",fontsize=10)
        ax.set_xlabel("销售渠道",fontsize=8)
        ax.set_ylabel("销售额（元）",fontsize=8)
        plt.xticks(fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        top_channel = channel_sales.idxmax()
        st.success(f"📌 **最有效的销售渠道：{top_channel}**，贡献 {channel_sales.max()/channel_sales.sum()*100:.1f}% 的销售额。")
        st.write("**数据解读：** 单一渠道成为营收支柱，其余渠道创收能力偏弱。")
        st.info("**经营决策：** 加大主力渠道推广预算，优化低效渠道运营模式或缩减投入。")

# 支付&新老客双饼
if pay_col or customer_type_col:
    col1, col2 = st.columns(2)
    with col1:
        if pay_col:
            df_clean["支付_清洗"] = clean_group_data(df_clean[pay_col])
            pay_sales = df_clean.groupby("支付_清洗")["销售额"].sum().sort_values(ascending=False)
            if not pay_sales.empty:
                fig, ax = plt.subplots(figsize=(3.2, 3), dpi=120)
                pay_sales.plot(kind="pie", autopct="%1.1f%%", ax=ax, wedgeprops=dict(width=0.7))
                ax.set_title("各支付方式销售额占比",fontsize=10)
                ax.set_ylabel("",fontsize=8)
                plt.tight_layout(pad=0.3)
                st.pyplot(fig, use_container_width=False)
                top_pay = pay_sales.index[0]
                st.success(f"📌 **最常用支付方式：{top_pay}**，占比 {pay_sales.iloc[0]/pay_sales.sum()*100:.1f}%。")
                st.write("**数据解读：** 用户支付习惯高度集中，主流支付占绝对份额。")
                st.info("**经营决策：** 优先保障主流支付通道稳定，保留小众支付满足多元化需求。")
    with col2:
        if customer_type_col:
            df_clean["顾客_清洗"] = clean_group_data(df_clean[customer_type_col])
            customer_sales = df_clean.groupby("顾客_清洗")["销售额"].sum()
            if not customer_sales.empty:
                fig, ax = plt.subplots(figsize=(3.2, 3), dpi=120)
                customer_sales.plot(kind="pie", autopct="%1.1f%%", ax=ax, wedgeprops=dict(width=0.7))
                ax.set_title("新顾客 vs 老顾客销售额占比",fontsize=10)
                ax.set_ylabel("",fontsize=8)
                plt.tight_layout(pad=0.3)
                st.pyplot(fig, use_container_width=False)
                top_cust = customer_sales.index[0]
                st.success(f"📌 **主要贡献顾客类型：{top_cust}**，占比 {customer_sales.max()/customer_sales.sum()*100:.1f}%。")
                st.write("**数据解读：** {top_cust}是营收核心来源。")
                st.info("**经营决策：** 老客占比高则完善会员储值；新客占比高则持续投放拉新活动。")

# ==================== 路径2：相关性与影响因素 ====================
st.divider()
st.subheader("路径2：相关性与影响因素分析")
numeric_cols = df_clean.select_dtypes(include=[np.number]).columns.tolist()
numeric_cols_for_corr = [c for c in numeric_cols if c not in ["地区_编码", "产品类别_编码"]]
if len(numeric_cols_for_corr) > 1:
    st.markdown("### 数值变量相关性热力图")
    corr_matrix = df_clean[numeric_cols_for_corr].corr()
    fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
    sns.heatmap(corr_matrix, annot=True, cmap="coolwarm", fmt=".2f", ax=ax, square=True, linewidths=0.5)
    ax.set_title("各数值变量相关性热力图", fontsize=9)
    plt.tight_layout(pad=0.3)
    st.pyplot(fig, use_container_width=False)
    if "销售额" in corr_matrix.columns:
        sales_corr = corr_matrix["销售额"].drop("销售额").sort_values(ascending=False)
        st.write("**与销售额相关性最高的因素：**")
        for var, corr_val in sales_corr.head(5).items():
            st.write(f"- {var}: {corr_val:.3f}")
        st.write("**数据解读：** 系数＞0正向拉动营收，系数＜0反向抑制营收，相关系数绝对值越大影响越强。")
        st.info("**经营决策：** 重点优化高正向相关指标、控制高负相关指标来提升销售额。")

# 价格-销量散点
if "价格" in df_clean.columns and "销量" in df_clean.columns:
    st.markdown("### 价格与销量关系分析")
    df_price_sales = df_clean[df_clean["价格"].notna() & df_clean["销量"].notna()].copy()
    if not df_price_sales.empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        ax.scatter(df_price_sales["价格"], df_price_sales["销量"], alpha=0.5, c='steelblue')
        ax.set_title("价格与销量散点图",fontsize=10)
        ax.set_xlabel("价格（元）",fontsize=8)
        ax.set_ylabel("销量",fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        st.write("**数据解读：** 多数样本随价格上涨销量下滑，符合需求规律，部分高价产品依旧保有销量。")
        st.info("**经营决策：** 常规产品低价走量、刚需稀缺产品维持高定价。")

# 折扣-销售额折线
if "折扣" in df_clean.columns:
    st.markdown("### 折扣力度与销售额关系")
    discount_sales = df_clean.groupby("折扣")["销售额"].sum()
    fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
    ax.plot(discount_sales.index, discount_sales.values, marker='o', linewidth=2, color='purple')
    ax.set_title("折扣力度与总销售额关系",fontsize=10)
    ax.set_xlabel("折扣",fontsize=8)
    ax.set_ylabel("总销售额（元）",fontsize=8)
    ax.grid(True, alpha=0.3)
    plt.tight_layout(pad=0.3)
    st.pyplot(fig, use_container_width=False)
    peak_disc = discount_sales.idxmax()
    st.write(f"**数据解读：** 折扣{peak_disc}对应销售额峰值，折扣过高/过低都会造成营收下滑。")
    st.info("**经营决策：** 日常促销固定选用最优{peak_disc}折扣，避免盲目低价亏损。")

# 年龄客单价柱状
if age_col:
    st.markdown("### 年龄对消费金额的影响")
    df_age_valid = df_clean[(df_clean["年龄_数值"] >= 0) & (df_clean["年龄_数值"] <= 120)]
    if not df_age_valid.empty:
        df_age_valid["年龄分组_细"] = pd.cut(df_age_valid["年龄_数值"], bins=10)
        age_amount = df_age_valid.groupby("年龄分组_细")["销售额"].mean()
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        ax.bar(range(len(age_amount)), age_amount.values, color='teal')
        ax.set_title("各年龄段的平均消费金额",fontsize=10)
        ax.set_xlabel("年龄段",fontsize=8)
        ax.set_ylabel("平均销售额（元）",fontsize=8)
        ax.set_xticks(range(len(age_amount)))
        ax.set_xticklabels([str(x) for x in age_amount.index], rotation=45, fontsize=7)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        st.write("**数据解读：** 不同年龄段单笔消费差距明显，部分群体客单价显著高于其他人群。")
        st.info("**经营决策：** 高客单价人群主推高端产品，低客单价人群主推平价刚需品。")

# ==================== 路径3：多模型对比【重点分层：数据指标解读→图像分析→结论决策】 ====================
st.divider()
st.subheader("路径3：多模型对比分析【重点分析模块】")
if "销售额" not in df_clean.columns:
    st.warning("未找到销售额列，跳过建模分析")
else:
    exclude_cols = ["销售额", "理论销售额", "地区_清洗", "产品_清洗", "类别_清洗", "销售_清洗",
                    "支付_清洗", "渠道_清洗", "口味", "年龄分组", "年龄分组_细", date_col,
                    "地区_编码", "产品类别_编码"]
    feature_candidates = [c for c in df_clean.columns if c not in exclude_cols and np.issubdtype(df_clean[c].dtype, np.number)]
    if len(feature_candidates) == 0:
        st.warning("无有效数值特征，跳过建模")
    else:
        model_data = df_clean[feature_candidates + ["销售额"]].dropna()
        if len(model_data) < 10:
            st.warning("有效建模数据不足10条，跳过模型训练")
        else:
            X = model_data[feature_candidates]
            y = model_data["销售额"]
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25, random_state=42)
            models = {
                "多元线性回归": LinearRegression(),
                "决策树回归": DecisionTreeRegressor(max_depth=5, random_state=42),
                "支持向量回归(SVR)": SVR(kernel="linear"),
                "随机森林回归": RandomForestRegressor(n_estimators=150, random_state=42)
            }
            model_results = {}
            model_preds = {}
            for name, model in models.items():
                model.fit(X_train, y_train)
                pred = model.predict(X_test)
                r2 = round(r2_score(y_test, pred), 4)
                mae = round(mean_absolute_error(y_test, pred), 2)
                rmse = round(np.sqrt(mean_squared_error(y_test, pred)), 2)
                model_results[name] = {"R²": r2, "MAE": mae, "RMSE": rmse}
                model_preds[name] = pred
            result_df = pd.DataFrame(model_results).T
            st.dataframe(result_df)
            best_model_name = max(model_results, key=lambda x: model_results[x]["R²"])
            best_results = model_results[best_model_name]
            worst_name = min(model_results, key=lambda x: model_results[x]["R²"])
            worst_results = model_results[worst_name]
            st.success(f"✅ 最优模型：{best_model_name}，R²={best_results['R²']}")

            # 1、表格指标数据解读
            st.markdown("##### ① 模型指标数据解读（从表格能看出）")
            st.write(f'''
1.R²：最优{best_model_name}(R²={best_results["R²"]})可解释{best_results["R²"]*100:.1f}%营收波动；{worst_name}(R²={worst_results["R²"]})拟合最差，解释能力不足。
2.MAE：平均预测误差，数值越小单样本预估偏差越小；RMSE：整体误差离散度，最优模型两项误差全模型最低。
''')
            # 2、落地决策
            st.markdown("##### ② 业务落地决策")
            st.info(f'''后续营收预测优先选用{best_model_name}；{worst_name}不适用本数据集，放弃作为预估工具；若最优R²偏低，需要新增营销、节假日等特征优化模型。''')

            # 真实vs预测+残差双图
            best_pred = model_preds[best_model_name]
            fig, axes = plt.subplots(1, 2, figsize=(10, 3), dpi=120)
            axes[0].scatter(y_test, best_pred, alpha=0.6, c='blue')
            axes[0].plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--', lw=2)
            axes[0].set_title("真实值 vs 预测值", fontsize=9)
            axes[0].set_xlabel("真实销售额（元）", fontsize=8)
            axes[0].set_ylabel("预测销售额（元）", fontsize=8)
            residuals = y_test - best_pred
            axes[1].scatter(best_pred, residuals, alpha=0.6, c='purple')
            axes[1].axhline(y=0, color='r', linestyle='--')
            axes[1].set_title("残差图", fontsize=9)
            axes[1].set_xlabel("预测销售额（元）", fontsize=8)
            axes[1].set_ylabel("残差", fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            # 图片解读+决策
            st.markdown("##### ③ 拟合图像数据分析")
            st.write("**数据解读：** 样本点越贴近红色对角线代表预测精准；残差均匀散乱在0轴线上下，无聚集偏移，模型无系统性偏差。")
            st.info("**模型优化决策：** 图像无异常偏移，当前模型可用；若残差集中一侧，需要清洗异常数据、补充特征。")

            # 特征重要性柱状
            final_model = models[best_model_name]
            if hasattr(final_model, 'feature_importances_'):
                importance = pd.Series(final_model.feature_importances_, index=feature_candidates)
            elif hasattr(final_model, 'coef_'):
                importance = pd.Series(np.abs(final_model.coef_), index=feature_candidates)
            else:
                importance = pd.Series([0]*len(feature_candidates), index=feature_candidates)
            importance = importance.sort_values(ascending=False)
            fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
            importance.head(10).plot(kind="bar", color="#2ca02c", ax=ax)
            ax.set_title("变量重要性排序（前10）",fontsize=10)
            ax.set_xlabel("特征变量",fontsize=8)
            ax.set_ylabel("重要性",fontsize=8)
            plt.xticks(rotation=45, ha="right", fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            top_feat = importance.index[0]
            st.write(f"**数据解读：** {top_feat}为影响销售额第一核心变量，重要度远超其余指标。")
            st.info(f"**经营决策：** 优先调控{top_feat}指标，是提升整体营收效率最高的手段。")

# ==================== 路径4：交互交叉分析 ====================
st.divider()
st.subheader("路径4：变量交互分析")
# 地区*品类
if region_col and category_col:
    st.markdown("### 地区 × 产品类别 交叉分析")
    cross_data = df_clean.groupby([region_col, category_col])["销售额"].sum().unstack(fill_value=0)
    if not cross_data.empty and len(cross_data) <= 20:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        cross_data.plot(kind="bar", ax=ax, stacked=False)
        ax.set_title("各地区不同产品类别销售额对比",fontsize=10)
        ax.set_xlabel("地区",fontsize=8)
        ax.set_ylabel("销售额（元）",fontsize=8)
        plt.xticks(rotation=45, ha="right", fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        st.dataframe(cross_data)
        st.write("**数据解读：** 不同地区热销品类完全不同，区域消费偏好差异化显著。")
        st.info("**经营决策：** 依据各地热销品类差异化铺货，高销品类增加区域库存。")

# 年龄*品类
if age_col and category_col:
    st.markdown("### 年龄 × 产品类别 交叉分析")
    df_age_valid = df_clean[(df_clean["年龄_数值"] >= 0) & (df_clean["年龄_数值"] <= 120)]
    if not df_age_valid.empty:
        df_age_valid["年龄分组"] = pd.cut(df_age_valid["年龄_数值"], bins=[0,25,35,45,60,120],
                                          labels=["18-25岁", "26-35岁", "36-45岁", "46-60岁", "60+岁"])
        cross_age = df_age_valid.groupby(["年龄分组", category_col])["销售额"].sum().unstack(fill_value=0)
        if not cross_age.empty:
            fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
            cross_age.plot(kind="bar", ax=ax)
            ax.set_title("各年龄层不同产品类别销售额对比",fontsize=10)
            ax.set_xlabel("年龄分组",fontsize=8)
            ax.set_ylabel("销售额（元）",fontsize=8)
            plt.xticks(fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            st.write("**数据解读：** 各年龄段偏好品类区分明显，同类产品在不同人群销量差异巨大。")
            st.info("**经营决策：** 按年龄段定向推荐对应品类产品，精准营销提升转化。")

# 渠道*支付
if channel_col and pay_col:
    st.markdown("### 销售渠道 × 支付方式 交叉分析")
    cross_channel_pay = pd.crosstab(df_clean[channel_col], df_clean[pay_col], values=df_clean["销售额"], aggfunc="sum", normalize="index") * 100
    if not cross_channel_pay.empty:
        fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
        cross_channel_pay.plot(kind="bar", ax=ax, stacked=True)
        ax.set_title("各渠道支付方式占比（%）",fontsize=10)
        ax.set_xlabel("销售渠道",fontsize=8)
        ax.set_ylabel("占比（%）",fontsize=8)
        plt.xticks(rotation=45, ha="right", fontsize=8)
        plt.tight_layout(pad=0.3)
        st.pyplot(fig, use_container_width=False)
        st.write("**数据解读：** 各渠道用户支付习惯不一致，A渠道主流支付在B渠道占比偏低。")
        st.info("**经营决策：** 各渠道配置自身用户偏好的主流支付方式，减少付款流失。")

# 价格*折扣
if "价格" in df_clean.columns and "折扣" in df_clean.columns:
    st.markdown("### 价格 × 折扣 交叉分析")
    df_price_valid = df_clean[df_clean["价格"].notna() & (df_clean["价格"] > 0)].copy()
    if not df_price_valid.empty:
        df_price_valid["价格区间"] = pd.cut(df_price_valid["价格"], bins=5)
        df_price_valid["折扣区间"] = pd.cut(df_price_valid["折扣"], bins=[0, 0.7, 0.85, 0.95, 1.0], labels=["低折扣", "中折扣", "高折扣", "无折扣"])
        cross_price = pd.pivot_table(df_price_valid, values="销售额", index="价格区间", columns="折扣区间", aggfunc="sum", fill_value=0)
        if not cross_price.empty:
            fig, ax = plt.subplots(figsize=(5, 3), dpi=120)
            cross_price.plot(kind="bar", ax=ax)
            ax.set_title("不同价格区间-折扣力度的销售额对比",fontsize=10)
            ax.set_xlabel("价格区间",fontsize=8)
            ax.set_ylabel("销售额（元）",fontsize=8)
            plt.xticks(rotation=45, ha="right", fontsize=8)
            plt.tight_layout(pad=0.3)
            st.pyplot(fig, use_container_width=False)
            st.write("**数据解读：** 高价商品依赖折扣拉动销量，低价商品对折扣敏感度极低。")
            st.info("**经营决策：** 高价产品搭配中高折扣促销，低价产品减少让利、维持原价走量。")
# ==================== 报告导出功能 ====================
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
# ==========新增开始==========
conclusion_items = locals().get("conclusion_items", ["数据清洗完成，基础指标正常"])
enterprise_advice = locals().get("enterprise_advice", "暂无专项经营优化建议")
top1 = locals().get("top1", "未识别关键影响指标")
top2 = locals().get("top2", "未识别次要影响指标")
# ==========新增结束==========
st.divider()
st.subheader("报告与数据导出")
report_content = f"""
企业销售数据分析与经营决策报告
================================

一、报告说明
本报告基于企业真实销售数据自动生成，已完成数据清洗、异常检测、建模分析，
所有结论均由数据驱动，可直接用于经营决策。

二、基础层：数据质量评估
- 原始数据行数：{original_rows}
- 清洗后有效数据：{final_rows}
- 检测到的脏数据：完全重复{dup_num}条，缺失值{miss_num}个，逻辑错误{zero_price_non_sales + zero_sales_non_amount}条
- 数据质量评级：良好

三、描述性分析核心结论
{chr(10).join(['- '+item for item in conclusion_items]) if conclusion_items else '- 无足够数据生成结论'}

四、企业综合经营建议
{enterprise_advice}

五、相关性分析结论
- 影响销售业绩的第一关键因素：{top1}
- 影响销售业绩的第二关键因素：{top2}
"""

if 'sales_corr' in locals() and not sales_corr.empty:
    report_content += f"\n- 与销售额相关性最高的变量：{sales_corr.index[0]}（相关系数{sales_corr.iloc[0]:.3f}）"

report_content += f"""
六、多模型对比结论
- 最优预测模型：{best_model_name if 'best_model_name' in locals() else '未训练'}
- 模型拟合精度 R²：{best_results['R²'] if 'best_results' in locals() else 0}
- 模型平均误差 MAE：{best_results['MAE'] if 'best_results' in locals() else 0}

七、变量交互分析结论
- 地区与产品类别存在明显的交叉影响
- 不同年龄层的产品偏好存在显著差异
- 不同渠道的主流支付方式各不相同

八、总结
本次分析通过全流程数据处理与智能建模，精准识别企业增长动力，
核心增长抓手为：{top1}、{top2}。

报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""

# 生成WORD报告：移除多余try捕获，优化BytesIO返回二进制流，解决下载docx损坏/空白报错
def create_word_report(content):
    doc = Document()
    title = doc.add_heading('企业销售数据分析与经营决策报告', 0)
    title.alignment = 1
    for line in content.split("\n"):
        if line.strip() == "":
            continue
        if line.startswith("="):
            continue
        if any(level in line for level in ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、"]):
            doc.add_heading(line.strip(), level=1)
        elif line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or line.startswith("4."):
            doc.add_paragraph(line.strip(), style='List Number')
        else:
            para = doc.add_paragraph(line.strip())
            for run in para.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    # 关键修改：返回字节数据，修复streamlit下载文件损坏
    return buffer.getvalue()

word_file = create_word_report(report_content)

col1, col2 = st.columns(2)
with col1:
    @st.cache_data
    def get_clean_data():
        return df_clean.to_csv(index=False, encoding="utf-8-sig")
    st.download_button("📥 下载清洗后数据（CSV）", data=get_clean_data(), file_name="企业清洗后销售数据.csv", mime="text/csv")
with col2:
    # 二进制数据直接填入data，修复下载报错
    st.download_button("📥 下载完整分析报告（WORD）", data=word_file, file_name="企业销售数据分析报告.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("### 清洗后数据预览")
st.dataframe(df_clean.head(20))

# 整体分析框架展示
st.divider()
with st.expander("📖 整体分析框架（设计思路）", expanded=False):
    st.markdown("""
    **一、整体分析框架**
    1. 基础层：数据清洗与质量评估
    2. 路径1：描述性分析（趋势/分布/对比）
    3. 路径2：相关性与影响因素分析
    4. 路径3：多模型对比分析
    5. 路径4：变量交互分析
    6. 报告导出
    """)

st.markdown("---")
st.markdown("**系统运行完成。如需重新分析，请刷新页面或重新上传数据。**")
